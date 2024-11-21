from rest_framework import viewsets, status
from rest_framework.decorators import api_view
from rest_framework.response import Response
from .models import Militar, Visitante, ViaturaAdministrativa, ViaturaOperacional, RegistroAcesso
from .serializers import (
    MilitarSerializer, 
    VisitanteSerializer, 
    ViaturaAdministrativaSerializer, 
    ViaturaOperacionalSerializer, 
    RegistroAcessoSerializer
)
from django.http import HttpResponse
from django.utils import timezone
from django.db.models import Q
from rest_framework.parsers import MultiPartParser, FormParser
from docx import Document
from datetime import datetime, timedelta
from docx.enum.table import WD_TABLE_ALIGNMENT
import io


# ViewSet para Militar
class MilitarViewSet(viewsets.ModelViewSet):
    queryset = Militar.objects.all()
    serializer_class = MilitarSerializer
    parser_classes = (MultiPartParser, FormParser) 

# ViewSet para Visitante
class VisitanteViewSet(viewsets.ModelViewSet):
    queryset = Visitante.objects.all()
    serializer_class = VisitanteSerializer
    parser_classes = (MultiPartParser, FormParser) 

# ViewSet para Viatura Administrativa
class ViaturaAdministrativaViewSet(viewsets.ModelViewSet):
    queryset = ViaturaAdministrativa.objects.all()
    serializer_class = ViaturaAdministrativaSerializer

# ViewSet para Viatura Operacional
class ViaturaOperacionalViewSet(viewsets.ModelViewSet):
    queryset = ViaturaOperacional.objects.all()
    serializer_class = ViaturaOperacionalSerializer

# ViewSet para Registro de Acesso
class RegistroAcessoViewSet(viewsets.ModelViewSet):
    queryset = RegistroAcesso.objects.all()
    serializer_class = RegistroAcessoSerializer

    # Sobrescrevendo create para lógica personalizada
    def create(self, request, *args, **kwargs):
        data = request.data
        tipo_acesso = data.get('tipo_acesso')

        # Validações adicionais para Viatura
        if tipo_acesso in ['VTR_ADM', 'VTR_OP']:
            if 'odometro' not in data or not data['odometro']:
                return Response(
                    {"error": "O campo 'odometro' é obrigatório para viaturas."},
                    status=status.HTTP_400_BAD_REQUEST
                )
            if 'motorista' not in data or not data['motorista']:
                return Response(
                    {"error": "O campo 'motorista' é obrigatório para viaturas."},
                    status=status.HTTP_400_BAD_REQUEST
                )

            # Validações para chefe de viatura em viaturas não administrativas
            if tipo_acesso == 'VTR_OP' and ('chefe_viatura' not in data or not data['chefe_viatura']):
                return Response(
                    {"error": "O campo 'chefe_viatura' é obrigatório para viaturas operacionais."},
                    status=status.HTTP_400_BAD_REQUEST
                )

        return super().create(request, *args, **kwargs)


@api_view(['POST'])
def registrar_acesso_qr_code(request):
    qr_code = request.data.get('qr_code')

    if not qr_code:
        return Response({'error': 'QR Code é obrigatório.'}, status=status.HTTP_400_BAD_REQUEST)

    tipo_acesso, codigo = qr_code.split(':')

    if(tipo_acesso == "MILITAR"):
        militar = Militar.objects.get(qr_code=codigo)
        if not militar:
            Response({'error': 'Código não reconhecido'}, status=status.HTTP_400_BAD_REQUEST)

        open_record = RegistroAcesso.objects.filter(
            militar=militar,
            data_hora_saida__isnull=True
        ).last()
        if open_record:
            militar.dentro=False
    
    elif(tipo_acesso == "VISITANTE"):
        visitante = Visitante.objects.get(qr_code=codigo)
        if not visitante:
            Response({'error': 'Código não reconhecido'}, status=status.HTTP_400_BAD_REQUEST)
        open_record = RegistroAcesso.objects.filter(
            visitante=visitante,
            data_hora_saida__isnull=True
        ).last()
        if open_record:
            visitante.dentro=False

    if open_record:
        open_record.data_hora_saida = timezone.now()
        open_record.save()
        return Response({'message': 'Saída registrada com sucesso.'}, status=status.HTTP_200_OK)
    
    registro = RegistroAcesso(
        tipo_acesso=tipo_acesso,
        data_hora_entrada=timezone.now(),
    )
    if tipo_acesso == 'MILITAR':
        militar.dentro = True
        registro.militar = militar
    elif tipo_acesso == 'VISITANTE':
        militar.dentro = True
        registro.visitante = visitante
    registro.save()
    return Response({'message': 'Entrada registrada com sucesso.'}, status=status.HTTP_201_CREATED)

@api_view(['POST'])
def registrar_acesso_viatura(request):
    
    data = request.data
    tipo_acesso = data['tipo_acesso']
    viatura_adm = data['viatura_administrativa']
    viatura_op = data['viatura_operacional']
    motorista = data['motorista']
    chefe_viatura = data['chefe_viatura']
    observacoes = data['observacoes']
    entrada_saida = data['entrada_saida']
    odometro = data['odometro']

    motorista = Militar.objects.get(id=motorista)
    chefe_viatura = Militar.objects.get(id= chefe_viatura)

    if(not motorista or not chefe_viatura):
        return Response({'error': 'Motorista ou chefe de viatura não encontrados.'}, status=status.HTTP_400_BAD_REQUEST)

    registro_motorista = RegistroAcesso(militar=motorista, tipo_acesso="MILITAR")
    registro_chefe_viatura = RegistroAcesso(militar=chefe_viatura, tipo_acesso="MILITAR")

    registro = RegistroAcesso(
        tipo_acesso = tipo_acesso,
        motorista = motorista,
        chefe_viatura = chefe_viatura,
        )
    
    if(entrada_saida == "ENTRADA"):
        registro.data_hora_entrada = timezone.now()
        registro_motorista.data_hora_entrada = timezone.now()
        registro_chefe_viatura.data_hora_entrada = timezone.now()
        
    elif(entrada_saida == "SAIDA"):
        registro.data_hora_saida = timezone.now()
        registro_motorista.data_hora_saida = timezone.now()
        registro_chefe_viatura.data_hora_saida = timezone.now()

    else:
        return Response({'error': 'Entrada ou saída devem ser especificadas.'}, status=status.HTTP_400_BAD_REQUEST)
    if viatura_adm:
        viatura_adm = ViaturaAdministrativa.objects.get(id = viatura_adm)
        if(entrada_saida == "ENTRADA"):
            viatura_adm.dentro = True
            viatura_adm.save()
        elif(entrada_saida == "SAIDA"):
            viatura_adm.dentro=False
            viatura_adm.save()
        registro.viatura_administrativa = viatura_adm
    else: 
        viatura_op = ViaturaOperacional.objects.get(id = viatura_op)
        if(entrada_saida == "ENTRADA"):
            viatura_op.dentro=True
            viatura_op.save()
        elif(entrada_saida == "SAIDA"):
            viatura_op.dentro=False
            viatura_op.save()
        registro.viatura_operacional = viatura_op
    registro.observacoes = observacoes
    registro.odometro = odometro
    registro.save()
    registro_motorista.save()
    registro_chefe_viatura.save()
    
    return Response({'message': f'{entrada_saida.lower()} registrada com sucesso.'}, status=status.HTTP_201_CREATED)



@api_view(['GET'])
def gerar_relatorio(request):
    # Obter parâmetros de filtro
    data_inicio = request.GET.get('data_inicio')
    data_fim = request.GET.get('data_fim')

    # Converter strings para objetos datetime
    try:
        if data_inicio:
            data_inicio = datetime.strptime(data_inicio, '%Y-%m-%d')
        if data_fim:
            data_fim = datetime.strptime(data_fim, '%Y-%m-%d')
    except ValueError:
        return Response({'error': 'Formato de data inválido. Use YYYY-MM-DD.'}, status=status.HTTP_400_BAD_REQUEST)

    # Filtrar registros de acesso com base nas datas
    registros = RegistroAcesso.objects.all()

    if data_inicio and data_fim:
        registros = registros.filter(
            Q(data_hora_entrada__gte=data_inicio) | Q(data_hora_entrada__lte=data_fim) |
            Q(data_hora_saida__gte=data_inicio) | Q(data_hora_saida__lte=data_fim)
        )
    elif data_inicio:
        registros = registros.filter(
            Q(data_hora_entrada__gte=data_inicio) | Q(data_hora_saida__gte=data_inicio)
        )
    elif data_fim:
        registros = registros.filter(
            Q(data_hora_entrada__lte=data_fim) | Q(data_hora_saida__lte=data_fim)
        )

    # Gerar o documento Word
    document = Document()
    document.add_heading('Relatório de Registros de Acesso', 0)

    # Adicionar detalhes dos filtros
    filtro_texto = 'Período: '
    if data_inicio:
        filtro_texto += f'De {data_inicio.strftime("%d/%m/%Y")} '
    if data_fim:
        filtro_texto += f'Até {data_fim.strftime("%d/%m/%Y")}'
    document.add_paragraph(filtro_texto)

    # Adicionar tabela com os registros
    table = document.add_table(rows=1, cols=5)
    table.style = 'Table Grid'  # Define um estilo para a tabela
    table.alignment = WD_TABLE_ALIGNMENT.CENTER  # Centraliza a tabela na página

    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Tipo de Acesso'
    hdr_cells[1].text = 'Identificação'
    hdr_cells[2].text = 'Data/Hora Entrada'
    hdr_cells[3].text = 'Data/Hora Saída'
    hdr_cells[4].text = 'Observações'

    # Permitir que as células do cabeçalho repitam em cada página
    table.rows[0].repeat_header = True

    # Definir a propriedade para permitir quebra de linha em todas as linhas
    for row in table.rows:
        row.height_rule = None  # Permite que a altura da linha se ajuste
        row.allow_break_across_pages = True  # Permite quebra entre páginas

    for registro in registros:
        row = table.add_row()
        row.allow_break_across_pages = True  # Permite quebra entre páginas
        row_cells = row.cells

        row_cells[0].text = registro.tipo_acesso
        if registro.militar:
            row_cells[1].text = f'{registro.militar.patente} {registro.militar.nome}'
        elif registro.visitante:
            row_cells[1].text = registro.visitante.nome
        elif registro.viatura_administrativa:
            row_cells[1].text = registro.viatura_administrativa.modelo
        elif registro.viatura_operacional:
            row_cells[1].text = registro.viatura_operacional.modelo
        else:
            row_cells[1].text = '---'

        row_cells[2].text = registro.data_hora_entrada.strftime('%d/%m/%Y %H:%M') if registro.data_hora_entrada else ''
        row_cells[3].text = registro.data_hora_saida.strftime('%d/%m/%Y %H:%M') if registro.data_hora_saida else ''
        row_cells[4].text = registro.observacoes or ''

    # Salvar o documento em um buffer de memória
    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)

    # Retornar o documento como uma resposta HTTP
    response = HttpResponse(
        buffer.getvalue(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    response['Content-Disposition'] = 'attachment; filename=relatorio_registros_acesso.docx'

    return response


@api_view(['GET'])
def gerar_relatorio_viaturas(request):
    # Obter parâmetros de filtro
    data_inicio = request.GET.get('data_inicio')
    data_fim = request.GET.get('data_fim')

    # Converter strings para objetos datetime
    try:
        if data_inicio:
            data_inicio = datetime.strptime(data_inicio, '%Y-%m-%d').date()
        if data_fim:
            data_fim = datetime.strptime(data_fim, '%Y-%m-%d').date()
    except ValueError:
        return Response({'error': 'Formato de data inválido. Use YYYY-MM-DD.'}, status=status.HTTP_400_BAD_REQUEST)

    # Verificar se as datas foram fornecidas
    if not data_inicio or not data_fim:
        return Response({'error': 'Por favor, forneça data de início e data de fim.'}, status=status.HTTP_400_BAD_REQUEST)

    # Garantir que data_inicio <= data_fim
    if data_inicio > data_fim:
        data_inicio, data_fim = data_fim, data_inicio

    # Gerar lista de datas no intervalo
    delta = data_fim - data_inicio
    lista_datas = [data_inicio + timedelta(days=i) for i in range(delta.days + 1)]

    # Gerar o documento Word
    document = Document()
    document.add_heading('Relatório de Viaturas', 0)

    # Adicionar detalhes dos filtros
    filtro_texto = f'Período: De {data_inicio.strftime("%d/%m/%Y")} Até {data_fim.strftime("%d/%m/%Y")}'
    document.add_paragraph(filtro_texto)

    for data in lista_datas:
        # Obter viaturas dentro e fora
        document.add_heading(f'Dia {data.strftime("%d/%m/%Y")}', level=1)

        # Obter viaturas que estavam dentro no dia
        viaturas_dentro_ids = RegistroAcesso.objects.filter(
            Q(viatura_administrativa__isnull=False) | Q(viatura_operacional__isnull=False),
            data_hora_entrada__date__lte=data,
        ).exclude(
            data_hora_saida__date__lt=data
        ).values_list('viatura_administrativa_id', 'viatura_operacional_id')

        viaturas_adm_dentro = ViaturaAdministrativa.objects.filter(id__in=[v[0] for v in viaturas_dentro_ids if v[0]])
        viaturas_op_dentro = ViaturaOperacional.objects.filter(id__in=[v[1] for v in viaturas_dentro_ids if v[1]])

        # Viaturas fora são aquelas não presentes em viaturas_dentro
        viaturas_adm_todas = ViaturaAdministrativa.objects.all()
        viaturas_op_todas = ViaturaOperacional.objects.all()

        viaturas_adm_fora = viaturas_adm_todas.exclude(id__in=viaturas_adm_dentro)
        viaturas_op_fora = viaturas_op_todas.exclude(id__in=viaturas_op_dentro)

        # Você pode ajustar a lógica acima se tiver um histórico de entradas e saídas.

        # Listar Viaturas Dentro
        document.add_heading('Viaturas Dentro da OM', level=2)
        if viaturas_adm_dentro.exists() or viaturas_op_dentro.exists():
            table = document.add_table(rows=1, cols=3)
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Tipo'
            hdr_cells[1].text = 'Modelo'
            hdr_cells[2].text = 'Identificação'

            # Viaturas Administrativas
            for viatura in viaturas_adm_dentro:
                row_cells = table.add_row().cells
                row_cells[0].text = 'Administrativa'
                row_cells[1].text = viatura.modelo
                row_cells[2].text = viatura.placa

            # Viaturas Operacionais
            for viatura in viaturas_op_dentro:
                row_cells = table.add_row().cells
                row_cells[0].text = 'Operacional'
                row_cells[1].text = viatura.modelo
                row_cells[2].text = viatura.eb_placa
        else:
            document.add_paragraph('Nenhuma viatura dentro.')

        # Listar Viaturas Fora
        document.add_heading('Viaturas Fora da OM', level=2)
        if viaturas_adm_fora.exists() or viaturas_op_fora.exists():
            table = document.add_table(rows=1, cols=3)
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Tipo'
            hdr_cells[1].text = 'Modelo'
            hdr_cells[2].text = 'Identificação'

            # Viaturas Administrativas
            for viatura in viaturas_adm_fora:
                row_cells = table.add_row().cells
                row_cells[0].text = 'Administrativa'
                row_cells[1].text = viatura.modelo
                row_cells[2].text = viatura.placa

            # Viaturas Operacionais
            for viatura in viaturas_op_fora:
                row_cells = table.add_row().cells
                row_cells[0].text = 'Operacional'
                row_cells[1].text = viatura.modelo
                row_cells[2].text = viatura.eb_placa
        else:
            document.add_paragraph('Nenhuma viatura fora.')

    # Salvar o documento em um buffer de memória
    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)

    # Retornar o documento como uma resposta HTTP
    response = HttpResponse(
        buffer.getvalue(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    response['Content-Disposition'] = 'attachment; filename=relatorio_viaturas.docx'

    return response